"""导出路由 - 支持 TXT / EPUB / DOCX"""
import io
import re
from urllib.parse import quote
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse


def clean_ai_artifacts(text: str) -> str:
    """清理 AI 生成的残留标记，用于导出前处理"""
    # 移除 [待补充] [此处待续] [略] 等方括号占位符
    text = re.sub(r'\[待[补充续写完]+\]', '', text)
    text = re.sub(r'\[此处[^\]]{0,20}\]', '', text)
    text = re.sub(r'\[略\]', '', text)
    text = re.sub(r'\[TODO\]', '', text, flags=re.IGNORECASE)

    # 移除 （此处描写xxx） 等圆括号占位符
    text = re.sub(r'（此处[^）]{0,30}）', '', text)
    text = re.sub(r'\(此处[^)]{0,30}\)', '', text)

    # 移除连续分隔线（3个以上的 - 或 * 或 =）
    text = re.sub(r'[-*═]{3,}\n?', '\n', text)
    text = re.sub(r'[=]{4,}\n?', '\n', text)

    # 移除 AI 自言自语（如 "以下是续写内容：" 等开头语）
    text = re.sub(r'^(以下是[^\n]{0,20}[：:]\s*\n)', '', text)
    text = re.sub(r'^(好的[，,][^\n]{0,20}[：:]\s*\n)', '', text)

    # 清理多余空行（超过2个连续空行合并为2个）
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()

router = APIRouter(prefix="/api/projects", tags=["export"])


async def _load_export_data(project_id: int) -> tuple[dict, list[dict]]:
    """加载导出所需的项目和章节数据（共享连接）"""
    from models.database import get_db_ctx
    async with get_db_ctx() as db:
        cursor = await db.execute("SELECT * FROM projects WHERE id=?", (project_id,))
        row = await cursor.fetchone()
        project = dict(row) if row else {}

        cursor = await db.execute(
            """SELECT chapter_number, title, content, word_count
               FROM chapters WHERE project_id=? AND content != ''
               ORDER BY chapter_number""",
            (project_id,),
        )
        chapters = [dict(r) for r in await cursor.fetchall()]

    return project, chapters


# ========== TXT 导出 ==========

@router.get("/{project_id}/export/txt")
async def export_txt(project_id: int):
    """导出纯文本 TXT"""
    project, chapters = await _load_export_data(project_id)
    if not chapters:
        raise HTTPException(404, "没有可导出的章节")

    lines = [f"《{project.get('name', '未命名')}》\n"]
    lines.append(f"类型：{project.get('genre', '未指定')}")
    lines.append(f"简介：{project.get('description', '')}\n")
    lines.append("=" * 40 + "\n")

    for ch in chapters:
        title = ch.get("title") or f"第{ch['chapter_number']}章"
        lines.append(f"\n{'=' * 20}\n{title}\n{'=' * 20}\n")
        lines.append(clean_ai_artifacts(ch["content"]))
        lines.append("")

    text = "\n".join(lines)
    filename = f"{project.get('name', 'novel')}.txt"
    encoded_filename = quote(filename)

    return StreamingResponse(
        io.BytesIO(text.encode("utf-8")),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )


# ========== DOCX 导出 ==========

@router.get("/{project_id}/export/docx")
async def export_docx(project_id: int):
    """导出 Word 文档 DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "DOCX 导出需要安装 python-docx")

    project, chapters = await _load_export_data(project_id)
    if not chapters:
        raise HTTPException(404, "没有可导出的章节")

    doc = Document()

    # 标题页
    title_para = doc.add_heading(level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run(f"《{project.get('name', '未命名')}》")

    if project.get("description"):
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = desc_para.add_run(project["description"])
        run.font.size = Pt(11)
        run.font.color.rgb = None  # 默认色

    doc.add_page_break()

    # 目录
    doc.add_heading("目录", level=1)
    for ch in chapters:
        title = ch.get("title") or f"第{ch['chapter_number']}章"
        doc.add_paragraph(title, style="List Number")
    doc.add_page_break()

    # 正文
    for ch in chapters:
        title = ch.get("title") or f"第{ch['chapter_number']}章"
        doc.add_heading(title, level=1)

        content = clean_ai_artifacts(ch["content"])
        # 按段落分割（空行分段）
        paragraphs = re.split(r'\n\s*\n', content)
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.line_spacing = 1.5

    # 保存到内存
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{project.get('name', 'novel')}.docx"
    encoded_filename = quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )


# ========== EPUB 导出 ==========

@router.get("/{project_id}/export/epub")
async def export_epub(project_id: int):
    """导出 EPUB 电子书"""
    try:
        from ebooklib import epub
    except ImportError:
        raise HTTPException(500, "EPUB 导出需要安装 ebooklib")

    project, chapters = await _load_export_data(project_id)
    if not chapters:
        raise HTTPException(404, "没有可导出的章节")

    book = epub.EpubBook()
    book.set_identifier(f"novel-{project_id}")
    book.set_title(project.get("name", "未命名"))
    book.set_language("zh")
    if project.get("description"):
        book.add_metadata("DC", "description", project["description"])

    # 样式
    style = """
    body { font-family: "PingFang SC", "Microsoft YaHei", sans-serif; line-height: 1.8; margin: 1em; }
    h1 { text-align: center; margin: 2em 0 1em; }
    p { text-indent: 2em; margin: 0.5em 0; }
    """
    css = epub.EpubItem(
        uid="style", file_name="style/default.css",
        media_type="text/css", content=style.encode("utf-8"),
    )
    book.add_item(css)

    # 封面页
    cover_html = f"""<html><head><link rel="stylesheet" href="style/default.css" /></head>
    <body><h1>《{project.get('name', '未命名')}》</h1>
    <p style="text-align:center;color:#666;">{project.get('genre', '')}</p>
    <p style="text-align:center;">{project.get('description', '')}</p>
    </body></html>"""
    cover = epub.EpubHtml(title="封面", file_name="cover.xhtml", lang="zh")
    cover.content = cover_html.encode("utf-8")
    cover.add_item(css)
    book.add_item(cover)

    # 章节
    epub_chapters = [cover]
    toc = []

    for ch in chapters:
        title = ch.get("title") or f"第{ch['chapter_number']}章"
        content = clean_ai_artifacts(ch["content"])

        # 按段落分割
        paragraphs = re.split(r'\n\s*\n', content)
        body_parts = []
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                body_parts.append(f"<p>{para_text}</p>")

        chapter_html = f"""<html><head>
        <link rel="stylesheet" href="style/default.css" /></head>
        <body><h1>{title}</h1>
        {"".join(body_parts)}
        </body></html>"""

        epub_ch = epub.EpubHtml(
            title=title,
            file_name=f"chapter_{ch['chapter_number']:04d}.xhtml",
            lang="zh",
        )
        epub_ch.content = chapter_html.encode("utf-8")
        epub_ch.add_item(css)
        book.add_item(epub_ch)
        epub_chapters.append(epub_ch)
        toc.append(epub_ch)

    # 目录和导航
    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + epub_chapters

    # 写入内存
    buf = io.BytesIO()
    epub.write_epub(buf, book)
    buf.seek(0)

    filename = f"{project.get('name', 'novel')}.epub"
    encoded_filename = quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/epub+zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )
